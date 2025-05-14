from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from langchain_community.chat_models import ChatOpenAI
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx2pdf import convert
from dotenv import load_dotenv
from pydantic import BaseModel, ValidationError
from typing import List, Optional
import tempfile, os, shutil, zipfile, smtplib, json, logging
from email.message import EmailMessage
import docx2txt
import fitz
from ats_audit import audit_resume

load_dotenv()
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

logging.basicConfig(level=logging.INFO)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

@app.get("/")
async def index():
    return {"status": "ok"}

class ResumeHeader(BaseModel):
    name: str
    email: str
    phone: str

class ResumeExperience(BaseModel):
    title: str
    company: str
    dates: str
    bullets: List[str]

class ResumeEducation(BaseModel):
    degree: str
    institution: str
    dates: str

class ResumeJSON(BaseModel):
    header: ResumeHeader
    summary: str
    experience: List[ResumeExperience]
    education: List[ResumeEducation]
    skills: List[str]

resume_prompt = PromptTemplate(
    input_variables=["resume_data", "job_desc", "tone"],
    template="""
SYSTEM:
You are “ResumeRevamp-GPT”, a former Fortune‑500 recruiter.
RULES:
1. Do not fabricate education, employers, dates, or metrics.
2. Use only facts that appear in either the candidate résumé (below) or the job description.
3. Each bullet must begin with an action verb, include a metric, and finish with a skill or tool.
4. Tone = {tone}. Acceptable values: formal | modern | confident | impactful.
5. ≤6 bullets per job, ≤15 words per bullet.
6. Return **valid JSON only** matching the schema – no markdown, no commentary.
SCHEMA:
{
  "header": {"name":"", "email":"", "phone":""},
  "summary": "",
  "experience": [
    {"title":"", "company":"", "dates":"", "bullets":[""]}
  ],
  "education": [
    {"degree":"", "institution":"", "dates":""}
  ],
  "skills": [""]
}
[JOB DESCRIPTION]
{job_desc}
[CANDIDATE RÉSUMÉ]
{resume_data}
"""
)

cover_prompt = PromptTemplate(
    input_variables=["resume_data", "job_desc", "tone", "company_name", "recipient_name"],
    template="""
SYSTEM: You are “CoverCraft‑GPT”, a career coach who crafts concise, story‑driven cover letters.
INSTRUCTIONS:
1. 250–300 words, four paragraphs (hook, alignment, value, close).
2. Mirror 5–7 keywords from the job description verbatim.
3. Mention ONE achievement from the résumé.
4. Close with a polite call‑to‑action.
Recipient = {recipient_name or "Hiring Manager"}
Company  = {company_name or "your organisation"}
Tone     = {tone}
[JOB DESCRIPTION]
{job_desc}
[RÉSUMÉ]
{resume_data}
"""
)

def style_doc(doc: DocxDocument) -> None:
    font = doc.styles["Normal"].font
    font.name = "Georgia"
    font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

def extract_text(file: UploadFile) -> str:
    suffix = file.filename.rsplit(".", 1)[-1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
        shutil.copyfileobj(file.file, tmp)
        path = tmp.name
    try:
        if suffix == "pdf":
            text = "\n".join(page.get_text() for page in fitz.open(path))
        elif suffix in {"docx", "doc"}:
            text = docx2txt.process(path)
        elif suffix == "txt":
            text = open(path, "r", encoding="utf-8").read()
        else:
            text = ""
    finally:
        os.unlink(path)
    return text

def build_resume_struct(resume_data: str, job_desc: str, tone: str = "confident") -> dict:
    llm = ChatOpenAI(model_name="gpt-4o", temperature=0.4)
    raw_json = LLMChain(llm=llm, prompt=resume_prompt).run({
        "resume_data": resume_data,
        "job_desc": job_desc,
        "tone": tone
    })
    try:
        parsed = json.loads(raw_json)
        ResumeJSON(**parsed)
        return parsed
    except (json.JSONDecodeError, ValidationError) as err:
        logging.error(f"Schema validation failed: {err}")
        return {"raw": raw_json}

def build_cover_letter(resume_data: str, job_desc: str, tone: str, company: str, recipient: str) -> str:
    if not recipient.strip(): recipient = "Hiring Manager"
    if not company.strip(): company = ""
    llm = ChatOpenAI(model_name="gpt-4o", temperature=0.5)
    return LLMChain(llm=llm, prompt=cover_prompt).run({
        "resume_data": resume_data,
        "job_desc": job_desc,
        "tone": tone,
        "company_name": company,
        "recipient_name": recipient
    })

def docx_from_struct(data: dict) -> DocxDocument:
    doc = DocxDocument()
    style_doc(doc)
    header = data.get("header", {})
    doc.add_paragraph(header.get("name", "")).bold = True
    contact_line = " | ".join(filter(None, [
        f"Email: {header.get('email', '')}" if header.get('email') else '',
        f"Phone: {header.get('phone', '')}" if header.get('phone') else ''
    ]))
    if contact_line:
        doc.add_paragraph(contact_line)
    if summary := data.get("summary"):
        doc.add_paragraph("")
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(summary.strip())
    if exps := data.get("experience"):
        doc.add_paragraph("")
        doc.add_heading("Experience", level=2)
        for role in exps:
            title_line = f"{role.get('title', '')} – {role.get('company', '')} ({role.get('dates', '')})"
            doc.add_paragraph(title_line).bold = True
            for b in role.get("bullets", []):
                if b.strip():
                    doc.add_paragraph(b.strip(), style="List Bullet")
    if edus := data.get("education"):
        doc.add_paragraph("")
        doc.add_heading("Education", level=2)
        for ed in edus:
            edu_line = f"{ed.get('degree', '')}, {ed.get('institution', '')} ({ed.get('dates', '')})"
            doc.add_paragraph(edu_line)
    if skills := data.get("skills"):
        doc.add_paragraph("")
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(skills))
    return doc

def email_resume_package(to_addr: str, from_addr: str, subject: str, body: str, attachments: List[str]):
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = from_addr
    msg["To"] = to_addr
    msg.set_content(body)
    for path in attachments:
        with open(path, "rb") as f:
            msg.add_attachment(f.read(), maintype="application", subtype="octet-stream", filename=os.path.basename(path))
    with smtplib.SMTP("smtp.gmail.com", 587) as server:
        server.starttls()
        server.login(os.getenv("SMTP_EMAIL"), os.getenv("SMTP_PASSWORD"))
        server.send_message(msg)

@app.post("/rewrite")
async def rewrite_resume(
    resume: UploadFile = File(...),
    tone: str = Form(...),
    job_desc: str = Form(...),
    company_name: str = Form(...),
    recipient_name: str = Form(...),
    sender_name: str = Form(...),
    sender_email: str = Form(...),
    sender_phone: str = Form(...),
    email_to: Optional[str] = Form(None),
    preview: str = Form("true")
):
    resume_text = extract_text(resume)
    resume_struct = build_resume_struct(resume_text, job_desc, tone)
    cover_text = build_cover_letter(resume_text, job_desc, tone, company_name, recipient_name)

    if preview.lower() == "true":
        audit_result = audit_resume(resume_struct, job_desc) if "raw" not in resume_struct else {}
        return JSONResponse({
            "resume_json": resume_struct,
            "cover_letter": cover_text.strip(),
            "ats_audit": audit_result
        })

    if "raw" in resume_struct:
        resume_doc = DocxDocument()
        resume_doc.add_paragraph(resume_struct["raw"])
    else:
        resume_doc = docx_from_struct(resume_struct)

    if "raw" not in resume_struct:
        first = resume_doc.paragraphs[0]
        if not first.text.strip():
            first.text = sender_name
        if not any("Email:" in p.text for p in resume_doc.paragraphs):
            resume_doc.paragraphs.insert(1, resume_doc.add_paragraph(f"Email: {sender_email} | Phone: {sender_phone}"))

    cover_doc = DocxDocument()
    style_doc(cover_doc)
    cover_doc.add_paragraph(sender_name).bold = True
    cover_doc.add_paragraph(f"Email: {sender_email} | Phone: {sender_phone}")
    cover_doc.add_paragraph("")
    for para in cover_text.split("\n"):
        if para.strip():
            cover_doc.add_paragraph(para.strip())

    with tempfile.TemporaryDirectory() as tmp:
        resume_path = os.path.join(tmp, "resume.docx")
        cover_path = os.path.join(tmp, "cover_letter.docx")
        zip_path = os.path.join(tmp, "resume_package.zip")

        resume_doc.save(resume_path)
        cover_doc.save(cover_path)
        convert(resume_path, resume_path.replace(".docx", ".pdf"))
        convert(cover_path, cover_path.replace(".docx", ".pdf"))

        if email_to:
            email_resume_package(
                to_addr=email_to,
                from_addr=sender_email,
                subject="Your Résumé + Cover Letter",
                body="Attached are your generated files.",
                attachments=[
                    resume_path,
                    cover_path,
                    resume_path.replace(".docx", ".pdf"),
                    cover_path.replace(".docx", ".pdf")
                ]
            )

        with zipfile.ZipFile(zip_path, "w") as zipf:
            zipf.write(resume_path, arcname="resume.docx")
            zipf.write(cover_path, arcname="cover_letter.docx")
            zipf.write(resume_path.replace(".docx", ".pdf"), arcname="resume.pdf")
            zipf.write(cover_path.replace(".docx", ".pdf"), arcname="cover_letter.pdf")

        return FileResponse(zip_path, media_type="application/zip", filename="resume_package.zip")
